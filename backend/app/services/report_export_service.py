"""
Report Export Service

Generates PDF and DOCX exports from HTML report content with institutional OppGrid branding.

Design System (matches oppgrid.com):
- Emerald accent: #10B981
- Dark navy: #0F172A (headings, table headers, OG badge)
- Slate palette: #334155 (body), #64748B (secondary), #94A3B8 (muted)
- Background tones: #F8FAFC (zebra), #F1F5F9 (borders), #E2E8F0 (dividers)
- Heading font: Georgia (serif) for report titles
- Body font: Helvetica/Arial (sans-serif) for everything else
"""

import html
import io
import json
import re
from datetime import datetime
from typing import Optional

from bs4 import BeautifulSoup
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ── Brand tokens ──────────────────────────────────────────────────────
EMERALD = "#10B981"
EMERALD_RGB = RGBColor(0x10, 0xB9, 0x81)
EMERALD_DARK = "#059669"
EMERALD_DARK_RGB = RGBColor(0x05, 0x96, 0x69)
EMERALD_LIGHT = "#ECFDF5"

NAVY = "#0F172A"
NAVY_RGB = RGBColor(0x0F, 0x17, 0x2A)

SLATE_800 = "#1E293B"
SLATE_800_RGB = RGBColor(0x1E, 0x29, 0x3B)
SLATE_700 = "#334155"
SLATE_700_RGB = RGBColor(0x33, 0x41, 0x55)
SLATE_500 = "#64748B"
SLATE_500_RGB = RGBColor(0x64, 0x74, 0x8B)
SLATE_400 = "#94A3B8"
SLATE_400_RGB = RGBColor(0x94, 0xA3, 0xB8)
SLATE_200 = "#E2E8F0"
SLATE_100 = "#F1F5F9"
SLATE_50 = "#F8FAFC"


# ── Economic snapshot helpers ─────────────────────────────────────────

_MACRO_LABELS: dict[str, str] = {
    "fed_funds_rate": "Fed Funds Rate",
    "inflation_rate": "CPI Index",
    "unemployment": "Unemployment Rate",
    "consumer_sentiment": "Consumer Sentiment",
    "mortgage_rate": "30-Yr Mortgage Rate",
    "gdp_growth": "GDP Growth",
}


def _fmt_indicator(value: float, units: str) -> str:
    if "%" in units or "percent" in units.lower():
        return f"{value:.2f}%"
    if "index" in units.lower() or "ratio" in units.lower():
        return f"{value:.1f}"
    return f"{value:.2f}"


def _fmt_revenue(value: float) -> str:
    if value >= 1e9:
        return f"${value / 1e9:.2f}B"
    if value >= 1e6:
        return f"${value / 1e6:.1f}M"
    if value >= 1e3:
        return f"${value / 1e3:.0f}K"
    return f"${value:.0f}"


def _fmt_employment(value: float) -> str:
    if value >= 1e6:
        return f"{value / 1e6:.2f}M"
    if value >= 1e3:
        return f"{value / 1e3:.1f}K"
    return f"{int(value):,}"


def _economic_snapshot_html(snapshot: dict) -> str:
    """Render economic snapshot data as an HTML appendix section for PDF export."""
    parts: list[str] = [
        '<div class="section" style="page-break-before: always;">',
        "<h1>Economic Intelligence</h1>",
        f'<p style="font-size:9pt; color:{SLATE_500}; margin-bottom:12px;">'
        "Live macro indicators, labor data, and public-comp benchmarks that grounded this report.</p>",
    ]

    macro: dict = snapshot.get("macro") or {}
    if macro:
        parts.append(f"<h2>Macroeconomic Environment</h2>")
        parts.append(
            f'<p style="font-size:8pt; color:{SLATE_400}; margin-bottom:6px;">'
            "Source: FRED / Federal Reserve</p>"
        )
        parts.append("<table>")
        parts.append("<tr><th>Indicator</th><th>Value</th><th>As Of</th></tr>")
        for key, indicator in macro.items():
            if not indicator:
                continue
            label = _MACRO_LABELS.get(key, key.replace("_", " ").title())
            value = indicator.get("value")
            units = indicator.get("units", "")
            date_str = indicator.get("date", "")
            if value is None:
                continue
            formatted = _fmt_indicator(float(value), units)
            date_display = ""
            if date_str:
                try:
                    d = datetime.fromisoformat(date_str[:10])
                    date_display = d.strftime("%b %Y")
                except Exception:
                    date_display = date_str[:7]
            parts.append(
                f"<tr><td>{html.escape(label)}</td>"
                f"<td><strong>{html.escape(formatted)}</strong></td>"
                f"<td>{html.escape(date_display)}</td></tr>"
            )
        parts.append("</table>")

    labor: dict | None = snapshot.get("labor")
    if labor:
        industry = html.escape(labor.get("industry_name", ""))
        source = html.escape(labor.get("source", "BLS"))
        naics = html.escape(str(labor.get("naics_code", "")))
        period = html.escape(str(labor.get("data_period", "")))
        emp = float(labor.get("total_employment", 0))
        yoy = float(labor.get("employment_change_yoy", 0))
        wage = float(labor.get("avg_weekly_wage", 0))
        estab = float(labor.get("establishment_count", 0))
        sign = "+" if yoy >= 0 else ""
        parts.append(f"<h2>Labor Market — {industry}</h2>")
        parts.append(
            f'<p style="font-size:8pt; color:{SLATE_400}; margin-bottom:6px;">{source}</p>'
        )
        parts.append("<table>")
        parts.append("<tr><th>Metric</th><th>Value</th></tr>")
        parts.append(f"<tr><td>Total Employment</td><td>{_fmt_employment(emp)} workers nationally</td></tr>")
        parts.append(f"<tr><td>YoY Employment Change</td><td>{sign}{yoy:.1f}%</td></tr>")
        parts.append(f"<tr><td>Avg Weekly Wage</td><td>${wage:,.0f} / week</td></tr>")
        if estab > 0:
            parts.append(f"<tr><td>Establishments</td><td>{_fmt_employment(estab)} businesses</td></tr>")
        parts.append("</table>")
        parts.append(
            f'<p style="font-size:8.5pt; color:{SLATE_400}; margin-top:4px;">'
            f"NAICS {naics} &middot; {period}</p>"
        )

    benchmarks: dict | None = snapshot.get("benchmarks")
    if benchmarks and benchmarks.get("public_comps"):
        avg_margin = float(benchmarks.get("avg_operating_margin", 0))
        avg_growth = benchmarks.get("avg_revenue_growth_3yr")
        source = html.escape(benchmarks.get("source", "SEC 10-K filings via sec-api.io"))
        comps: list[dict] = benchmarks["public_comps"]
        parts.append("<h2>Public-Comp Benchmarks</h2>")
        parts.append(
            f'<p style="font-size:8pt; color:{SLATE_400}; margin-bottom:6px;">{source}</p>'
        )
        parts.append("<table>")
        parts.append("<tr><th>Metric</th><th>Value</th></tr>")
        parts.append(f"<tr><td>Avg Operating Margin</td><td><strong>{avg_margin * 100:.1f}%</strong></td></tr>")
        if avg_growth is not None:
            g_sign = "+" if float(avg_growth) >= 0 else ""
            parts.append(f"<tr><td>3-Yr Revenue CAGR</td><td>{g_sign}{float(avg_growth) * 100:.1f}%</td></tr>")
        parts.append("</table>")
        parts.append("<h3>Comparable Companies</h3>")
        parts.append("<table>")
        parts.append("<tr><th>Ticker</th><th>Company</th><th>FY</th><th>Revenue</th><th>Op. Income</th><th>Margin</th></tr>")
        for comp in comps:
            ticker = html.escape(str(comp.get("ticker", "")))
            company = html.escape(str(comp.get("company_name", "")))
            fy = comp.get("fiscal_year", "")
            rev = float(comp.get("revenue", 0))
            op_inc = float(comp.get("operating_income", 0))
            margin = comp.get("operating_margin")
            margin_str = f"{float(margin) * 100:.1f}%" if margin is not None else "&mdash;"
            parts.append(
                f"<tr>"
                f"<td><strong>{ticker}</strong></td>"
                f"<td>{company}</td>"
                f"<td>{fy}</td>"
                f"<td>{_fmt_revenue(rev)}</td>"
                f"<td>{_fmt_revenue(op_inc)}</td>"
                f"<td>{margin_str}</td>"
                f"</tr>"
            )
        parts.append("</table>")

    parts.append("</div>")
    return "\n".join(parts)


def _add_economic_snapshot_docx(doc: Document, snapshot: dict) -> None:
    """Append an Economic Intelligence appendix to an open DOCX document."""

    def _label_row(table, label: str, value: str, zebra: bool = False) -> None:
        row = table.add_row()
        for cell, text in zip(row.cells, [label, value]):
            cell.text = text
            if zebra:
                _set_cell_shading(cell, SLATE_50)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = SLATE_700_RGB
                    run.font.name = "Arial"

    page_break = doc.add_paragraph()
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement
    run = page_break.add_run()
    br = OxmlElement("w:br")
    br.set(_qn("w:type"), "page")
    run._r.append(br)

    _add_heading(doc, "Economic Intelligence", level=1)
    intro = doc.add_paragraph()
    run = intro.add_run(
        "Live macro indicators, labor data, and public-comp benchmarks that grounded this report."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = SLATE_500_RGB
    run.font.name = "Arial"
    intro.paragraph_format.space_after = Pt(10)

    macro: dict = snapshot.get("macro") or {}
    if macro:
        _add_heading(doc, "Macroeconomic Environment", level=2)
        src_para = doc.add_paragraph()
        src_run = src_para.add_run("Source: FRED / Federal Reserve")
        src_run.font.size = Pt(8)
        src_run.font.color.rgb = SLATE_400_RGB
        src_run.font.name = "Arial"
        src_para.paragraph_format.space_after = Pt(4)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0]
        for cell, text in zip(hdr.cells, ["Indicator", "Value", "As Of"]):
            cell.text = text
            _set_cell_shading(cell, NAVY)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.name = "Arial"

        for i, (key, indicator) in enumerate(macro.items()):
            if not indicator:
                continue
            value = indicator.get("value")
            if value is None:
                continue
            label = _MACRO_LABELS.get(key, key.replace("_", " ").title())
            units = indicator.get("units", "")
            formatted = _fmt_indicator(float(value), units)
            date_str = indicator.get("date", "")
            date_display = ""
            if date_str:
                try:
                    d = datetime.fromisoformat(date_str[:10])
                    date_display = d.strftime("%b %Y")
                except Exception:
                    date_display = date_str[:7]
            row = table.add_row()
            values = [label, formatted, date_display]
            for j, (cell, text) in enumerate(zip(row.cells, values)):
                cell.text = text
                if i % 2 == 1:
                    _set_cell_shading(cell, SLATE_50)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9.5)
                        run.font.color.rgb = SLATE_700_RGB
                        run.font.name = "Arial"
                        if j == 1:
                            run.bold = True
        doc.add_paragraph()

    labor: dict | None = snapshot.get("labor")
    if labor:
        industry = labor.get("industry_name", "")
        source = labor.get("source", "BLS")
        naics = str(labor.get("naics_code", ""))
        period = str(labor.get("data_period", ""))
        emp = float(labor.get("total_employment", 0))
        yoy = float(labor.get("employment_change_yoy", 0))
        wage = float(labor.get("avg_weekly_wage", 0))
        estab = float(labor.get("establishment_count", 0))
        sign = "+" if yoy >= 0 else ""

        _add_heading(doc, f"Labor Market — {industry}", level=2)
        src_para = doc.add_paragraph()
        src_run = src_para.add_run(source)
        src_run.font.size = Pt(8)
        src_run.font.color.rgb = SLATE_400_RGB
        src_run.font.name = "Arial"
        src_para.paragraph_format.space_after = Pt(4)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0]
        for cell, text in zip(hdr.cells, ["Metric", "Value"]):
            cell.text = text
            _set_cell_shading(cell, NAVY)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.name = "Arial"

        rows = [
            ("Total Employment", f"{_fmt_employment(emp)} workers nationally"),
            ("YoY Employment Change", f"{sign}{yoy:.1f}%"),
            ("Avg Weekly Wage", f"${wage:,.0f} / week"),
        ]
        if estab > 0:
            rows.append(("Establishments", f"{_fmt_employment(estab)} businesses"))
        for i, (lbl, val) in enumerate(rows):
            _label_row(table, lbl, val, zebra=(i % 2 == 1))

        naics_para = doc.add_paragraph()
        naics_run = naics_para.add_run(f"NAICS {naics} · {period}")
        naics_run.font.size = Pt(8)
        naics_run.font.color.rgb = SLATE_400_RGB
        naics_run.font.name = "Arial"
        naics_para.paragraph_format.space_after = Pt(8)

    benchmarks: dict | None = snapshot.get("benchmarks")
    if benchmarks and benchmarks.get("public_comps"):
        avg_margin = float(benchmarks.get("avg_operating_margin", 0))
        avg_growth = benchmarks.get("avg_revenue_growth_3yr")
        source = benchmarks.get("source", "SEC 10-K filings via sec-api.io")
        comps: list[dict] = benchmarks["public_comps"]

        _add_heading(doc, "Public-Comp Benchmarks", level=2)
        src_para = doc.add_paragraph()
        src_run = src_para.add_run(source)
        src_run.font.size = Pt(8)
        src_run.font.color.rgb = SLATE_400_RGB
        src_run.font.name = "Arial"
        src_para.paragraph_format.space_after = Pt(4)

        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = "Table Grid"
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = summary_table.rows[0]
        for cell, text in zip(hdr.cells, ["Metric", "Value"]):
            cell.text = text
            _set_cell_shading(cell, NAVY)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.name = "Arial"
        _label_row(summary_table, "Avg Operating Margin", f"{avg_margin * 100:.1f}%", zebra=False)
        if avg_growth is not None:
            g_sign = "+" if float(avg_growth) >= 0 else ""
            _label_row(summary_table, "3-Yr Revenue CAGR", f"{g_sign}{float(avg_growth) * 100:.1f}%", zebra=True)
        doc.add_paragraph()

        _add_heading(doc, "Comparable Companies", level=3)
        comp_table = doc.add_table(rows=1, cols=6)
        comp_table.style = "Table Grid"
        comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = comp_table.rows[0]
        for cell, text in zip(hdr.cells, ["Ticker", "Company", "FY", "Revenue", "Op. Income", "Margin"]):
            cell.text = text
            _set_cell_shading(cell, NAVY)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.name = "Arial"

        for i, comp in enumerate(comps):
            ticker = str(comp.get("ticker", ""))
            company = str(comp.get("company_name", ""))
            fy = str(comp.get("fiscal_year", ""))
            rev = float(comp.get("revenue", 0))
            op_inc = float(comp.get("operating_income", 0))
            margin = comp.get("operating_margin")
            margin_str = f"{float(margin) * 100:.1f}%" if margin is not None else "—"
            row = comp_table.add_row()
            values = [ticker, company, fy, _fmt_revenue(rev), _fmt_revenue(op_inc), margin_str]
            for j, (cell, text) in enumerate(zip(row.cells, values)):
                cell.text = text
                if i % 2 == 1:
                    _set_cell_shading(cell, SLATE_50)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = SLATE_700_RGB
                        run.font.name = "Arial"
                        if j == 0:
                            run.bold = True
        doc.add_paragraph()


# ── PDF template ──────────────────────────────────────────────────────

def _branded_html_wrapper(content: str, title: str, report_type: str, generated_at: str) -> str:
    """Wrap report HTML content with institutional OppGrid-branded PDF template."""
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  @page {{
    size: letter;
    margin: 0.75in 0.85in 0.85in 0.85in;
    @frame footer {{
      -pdf-frame-content: page-footer;
      bottom: 0.3in;
      margin-left: 0.85in;
      margin-right: 0.85in;
      height: 0.4in;
    }}
  }}
  body {{
    font-family: Helvetica, Arial, sans-serif;
    font-size: 10.5pt;
    line-height: 1.65;
    color: {SLATE_700};
  }}

  /* ── Accent bar ── */
  .accent-bar {{
    background-color: {EMERALD};
    height: 5px;
    margin: -0.75in -0.85in 0 -0.85in;
  }}

  /* ── Masthead ── */
  .masthead {{
    padding: 20px 0 16px 0;
    border-bottom: 1px solid {SLATE_200};
    margin-bottom: 20px;
  }}
  .masthead .wordmark {{
    font-family: Helvetica, Arial, sans-serif;
    font-size: 17pt;
    font-weight: bold;
    color: {NAVY};
    letter-spacing: -0.3px;
    margin: 0;
    padding: 0;
    line-height: 1.2;
  }}
  .masthead .tagline {{
    font-family: Helvetica, Arial, sans-serif;
    font-size: 8pt;
    color: {SLATE_500};
    letter-spacing: 1.5px;
    text-transform: uppercase;
    margin: 0;
    padding: 0;
  }}

  /* ── Title block ── */
  .title-block {{
    margin-bottom: 8px;
  }}
  .report-type-badge {{
    font-family: Helvetica, Arial, sans-serif;
    font-size: 8.5pt;
    font-weight: bold;
    color: {EMERALD};
    letter-spacing: 2px;
    text-transform: uppercase;
    margin-bottom: 8px;
  }}
  .report-title {{
    font-family: Georgia, 'Times New Roman', serif;
    font-size: 24pt;
    font-weight: normal;
    color: {NAVY};
    margin: 0 0 4px 0;
    line-height: 1.2;
    letter-spacing: -0.3px;
  }}

  /* ── Metadata row ── */
  .meta-row {{
    margin-top: 16px;
    padding: 10px 0;
    border-top: 1px solid {SLATE_100};
    border-bottom: 1px solid {SLATE_100};
  }}
  .meta-row table {{
    width: 100%;
    border-collapse: collapse;
    border: none;
  }}
  .meta-row td {{
    border: none;
    padding: 0 16px 0 0;
    vertical-align: top;
    background: transparent;
  }}
  .meta-label {{
    font-family: Helvetica, Arial, sans-serif;
    font-size: 7.5pt;
    color: {SLATE_400};
    text-transform: uppercase;
    letter-spacing: 1px;
  }}
  .meta-value {{
    font-family: Helvetica, Arial, sans-serif;
    font-size: 10pt;
    color: {SLATE_700};
    margin-top: 2px;
  }}

  /* ── Section headings ── */
  h1 {{
    font-family: Georgia, 'Times New Roman', serif;
    font-size: 15pt;
    font-weight: normal;
    color: {NAVY};
    border-bottom: 2.5px solid {EMERALD};
    padding-bottom: 5px;
    margin-top: 28px;
    margin-bottom: 12px;
  }}
  h2 {{
    font-family: Helvetica, Arial, sans-serif;
    font-size: 12pt;
    font-weight: bold;
    color: {NAVY};
    margin-top: 20px;
    margin-bottom: 8px;
  }}
  h3 {{
    font-family: Helvetica, Arial, sans-serif;
    font-size: 11pt;
    font-weight: bold;
    color: {SLATE_800};
    margin-top: 16px;
    margin-bottom: 6px;
  }}

  /* ── Body ── */
  p {{ margin: 6px 0; }}
  ul, ol {{ margin: 6px 0; padding-left: 22px; }}
  li {{ margin-bottom: 3px; }}
  strong, b {{ color: {NAVY}; }}

  /* ── Tables ── */
  table {{
    width: 100%;
    border-collapse: collapse;
    margin: 12px 0;
    font-size: 9.5pt;
  }}
  th {{
    background-color: {NAVY};
    color: white;
    font-weight: 500;
    padding: 7px 10px;
    text-align: left;
    border: none;
  }}
  td {{
    padding: 6px 10px;
    border-bottom: 1px solid {SLATE_100};
    color: {SLATE_700};
  }}
  tr:nth-child(even) td {{
    background-color: {SLATE_50};
  }}

  /* ── Callout / insight blocks ── */
  blockquote {{
    border-left: 3px solid {EMERALD};
    margin: 12px 0;
    padding: 8px 16px;
    background-color: {EMERALD_LIGHT};
    font-style: italic;
    color: {SLATE_700};
  }}

  .section {{
    page-break-inside: avoid;
  }}
</style>
</head>
<body>
  <!-- Accent bar -->
  <div class="accent-bar"></div>

  <!-- Masthead -->
  <div class="masthead">
    <span class="wordmark">OppGrid</span><br/>
    <span class="tagline">Opportunity Intelligence</span>
  </div>

  <!-- Title block -->
  <div class="title-block">
    <div class="report-type-badge">{report_type}</div>
    <div class="report-title">{title}</div>
  </div>

  <!-- Metadata row -->
  <div class="meta-row">
    <table>
      <tr>
        <td>
          <div class="meta-label">Generated</div>
          <div class="meta-value">{generated_at}</div>
        </td>
        <td>
          <div class="meta-label">Platform</div>
          <div class="meta-value">oppgrid.com</div>
        </td>
      </tr>
    </table>
  </div>

  <!-- Report content -->
  {content}

  <!-- Footer -->
  <div id="page-footer" style="font-size: 7.5pt; color: {SLATE_400};">
    <table style="width: 100%; border: none; margin: 0; padding: 0;">
      <tr>
        <td style="border: none; padding: 0; text-align: left; background: transparent;">Confidential &mdash; Prepared by OppGrid AI</td>
        <td style="border: none; padding: 0; text-align: right; background: transparent;">oppgrid.com</td>
      </tr>
    </table>
  </div>
</body>
</html>"""


# ── Plain-text detection ──────────────────────────────────────────────

_HTML_TAG_RE = re.compile(r"<[a-zA-Z][^>]*>|</[a-zA-Z]+>", re.DOTALL)
_BORDER_ONLY_RE = re.compile(r"^[\s═─━╔╗╚╝║╠╣╦╩╪╫▀▄█▌▐░▒▓\-=_~*+|]+$")


def _ensure_html(content: str) -> str:
    """Convert plain-text report content to basic HTML if needed."""
    if _HTML_TAG_RE.search(content):
        return content

    html_parts: list[str] = []
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if _BORDER_ONLY_RE.match(stripped):
            continue
        html_parts.append(f"<p>{html.escape(stripped)}</p>")
    return "\n".join(html_parts)


# ── PDF generation ────────────────────────────────────────────────────

def generate_pdf(
    html_content: str,
    title: str = "OppGrid Report",
    report_type: str = "Report",
    generated_at: Optional[str] = None,
    economic_snapshot: Optional[dict] = None,
) -> bytes:
    """Convert HTML report content to an institutionally branded PDF."""
    if not generated_at:
        generated_at = datetime.utcnow().strftime("%B %d, %Y")

    html_content = _ensure_html(html_content)
    if economic_snapshot:
        html_content = html_content + "\n" + _economic_snapshot_html(economic_snapshot)
    full_html = _branded_html_wrapper(html_content, title, report_type, generated_at)

    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(full_html, dest=buffer, encoding="utf-8")

    if pisa_status.err:
        raise RuntimeError(f"PDF generation failed with {pisa_status.err} errors")

    buffer.seek(0)
    return buffer.read()


# ── DOCX helpers ──────────────────────────────────────────────────────

def _strip_tags(html_str: str) -> str:
    """Remove HTML tags, returning plain text."""
    return re.sub(r"<[^>]+>", "", html_str).strip()


def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell via XML."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn("w:shd"), {
        qn("w:fill"): hex_color.lstrip("#"),
        qn("w:val"): "clear",
    })
    tc_pr.append(shading)


def _add_para_border(paragraph, side: str, sz: str, color: str):
    """Add a border to a specific side of a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    border = pBdr.makeelement(qn(f"w:{side}"), {
        qn("w:val"): "single",
        qn("w:sz"): sz,
        qn("w:space"): "1",
        qn("w:color"): color.lstrip("#"),
    })
    pBdr.append(border)
    pPr.append(pBdr)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = NAVY_RGB
            run.font.size = Pt(16)
            run.font.name = "Georgia"
        elif level == 2:
            run.font.color.rgb = NAVY_RGB
            run.font.size = Pt(13)
            run.font.name = "Arial"
        else:
            run.font.color.rgb = SLATE_800_RGB
            run.font.size = Pt(11)
            run.font.name = "Arial"


def _add_paragraph(doc: Document, text: str, bold: bool = False):
    """Add a styled body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = SLATE_700_RGB
    run.font.name = "Arial"
    run.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p


# ── DOCX generation ──────────────────────────────────────────────────

def generate_docx(
    html_content: str,
    title: str = "OppGrid Report",
    report_type: str = "Report",
    generated_at: Optional[str] = None,
    economic_snapshot: Optional[dict] = None,
) -> bytes:
    """Convert HTML report content to an institutionally branded DOCX file."""
    if not generated_at:
        generated_at = datetime.utcnow().strftime("%B %d, %Y")

    html_content = _ensure_html(html_content)
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Accent line (emerald top rule) ──
    accent = doc.add_paragraph()
    accent.paragraph_format.space_after = Pt(16)
    _add_para_border(accent, "top", "24", EMERALD)

    # ── Wordmark + tagline (single paragraph to prevent Word separator) ──
    masthead = doc.add_paragraph()
    masthead.paragraph_format.space_before = Pt(0)
    masthead.paragraph_format.space_after = Pt(14)
    run = masthead.add_run("OppGrid")
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY_RGB
    run.font.name = "Arial"
    run.bold = True
    masthead.add_run("\n")
    run = masthead.add_run("OPPORTUNITY INTELLIGENCE")
    run.font.size = Pt(8)
    run.font.color.rgb = SLATE_500_RGB
    run.font.name = "Arial"
    run.bold = False

    # ── Divider ──
    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(14)
    _add_para_border(divider, "bottom", "4", SLATE_200)

    # ── Report type badge ──
    badge = doc.add_paragraph()
    badge.paragraph_format.space_after = Pt(4)
    run = badge.add_run(report_type.upper())
    run.font.size = Pt(8.5)
    run.font.color.rgb = EMERALD_RGB
    run.font.name = "Arial"
    run.bold = True

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(4)
    run = title_para.add_run(title)
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY_RGB
    run.font.name = "Georgia"

    # ── Metadata ──
    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(8)
    meta.paragraph_format.space_after = Pt(16)
    run = meta.add_run(f"{report_type}  |  Generated {generated_at}  |  oppgrid.com")
    run.font.size = Pt(9)
    run.font.color.rgb = SLATE_400_RGB
    run.font.name = "Arial"

    # ── Divider ──
    divider2 = doc.add_paragraph()
    divider2.paragraph_format.space_after = Pt(12)
    _add_para_border(divider2, "bottom", "4", SLATE_100)

    # ── Footer ──
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = footer_para.add_run("Confidential — Prepared by OppGrid AI")
        run.font.size = Pt(7.5)
        run.font.color.rgb = SLATE_400_RGB
        run.font.name = "Arial"

    # ── Parse HTML body ──
    soup = BeautifulSoup(html_content, "html.parser")
    for element in soup.children:
        _process_element(doc, element)

    # ── Economic Intelligence appendix ──
    if economic_snapshot:
        _add_economic_snapshot_docx(doc, economic_snapshot)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ── HTML-to-DOCX element processing ──────────────────────────────────

def _process_element(doc: Document, element):
    """Recursively process an HTML element into Word document elements."""
    if isinstance(element, str):
        text = element.strip()
        if text:
            _add_paragraph(doc, text)
        return

    tag = element.name
    if tag is None:
        return

    if tag in ("h1",):
        _add_heading(doc, element.get_text(strip=True), level=1)
    elif tag in ("h2",):
        _add_heading(doc, element.get_text(strip=True), level=2)
    elif tag in ("h3", "h4", "h5", "h6"):
        _add_heading(doc, element.get_text(strip=True), level=3)
    elif tag == "p":
        text = element.get_text(strip=True)
        if text:
            p = doc.add_paragraph()
            _build_inline_runs(p, element)
            p.paragraph_format.space_after = Pt(4)
    elif tag in ("ul", "ol"):
        for li in element.find_all("li", recursive=False):
            text = li.get_text(strip=True)
            if text:
                p = doc.add_paragraph(style="List Bullet" if tag == "ul" else "List Number")
                run = p.add_run(text)
                run.font.size = Pt(10.5)
                run.font.color.rgb = SLATE_700_RGB
                run.font.name = "Arial"
    elif tag == "table":
        _process_table(doc, element)
    elif tag == "blockquote":
        text = element.get_text(strip=True)
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            _add_para_border(p, "left", "12", EMERALD)
            run = p.add_run(text)
            run.font.size = Pt(10.5)
            run.font.color.rgb = SLATE_700_RGB
            run.font.name = "Arial"
            run.italic = True
    elif tag in ("div", "section", "article", "main", "span", "strong", "em", "b", "i"):
        for child in element.children:
            _process_element(doc, child)
    elif tag == "br":
        doc.add_paragraph()
    elif tag == "hr":
        divider = doc.add_paragraph()
        divider.paragraph_format.space_before = Pt(8)
        divider.paragraph_format.space_after = Pt(8)
    else:
        text = element.get_text(strip=True)
        if text:
            _add_paragraph(doc, text)


def _build_inline_runs(paragraph, element):
    """Build runs for inline elements (bold, italic, links) within a paragraph."""
    for child in element.children:
        if isinstance(child, str):
            text = child
            if text:
                run = paragraph.add_run(text)
                run.font.size = Pt(10.5)
                run.font.color.rgb = SLATE_700_RGB
                run.font.name = "Arial"
        elif child.name in ("strong", "b"):
            run = paragraph.add_run(child.get_text())
            run.font.size = Pt(10.5)
            run.font.color.rgb = NAVY_RGB
            run.font.name = "Arial"
            run.bold = True
        elif child.name in ("em", "i"):
            run = paragraph.add_run(child.get_text())
            run.font.size = Pt(10.5)
            run.font.color.rgb = SLATE_500_RGB
            run.font.name = "Arial"
            run.italic = True
        elif child.name == "a":
            run = paragraph.add_run(child.get_text())
            run.font.size = Pt(10.5)
            run.font.color.rgb = EMERALD_DARK_RGB
            run.font.name = "Arial"
            run.underline = True
        elif child.name == "br":
            paragraph.add_run("\n")
        else:
            run = paragraph.add_run(child.get_text())
            run.font.size = Pt(10.5)
            run.font.color.rgb = SLATE_700_RGB
            run.font.name = "Arial"


def _process_table(doc: Document, table_elem):
    """Convert an HTML table to a branded Word table with dark header and zebra rows."""
    rows_data = []
    header_row = None

    thead = table_elem.find("thead")
    if thead:
        tr = thead.find("tr")
        if tr:
            header_row = [th.get_text(strip=True) for th in tr.find_all(["th", "td"])]

    tbody = table_elem.find("tbody")
    body_rows = tbody.find_all("tr") if tbody else table_elem.find_all("tr")

    for tr in body_rows:
        cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
        if cells:
            if not header_row and tr.find("th"):
                header_row = cells
            else:
                rows_data.append(cells)

    if not rows_data and not header_row:
        return

    num_cols = max(
        len(header_row) if header_row else 0,
        max((len(r) for r in rows_data), default=0),
    )
    if num_cols == 0:
        return

    num_rows = len(rows_data) + (1 if header_row else 0)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    row_idx = 0

    # ── Header row: dark navy background ──
    if header_row:
        for col_idx, text in enumerate(header_row[:num_cols]):
            cell = table.cell(0, col_idx)
            cell.text = text
            _set_cell_shading(cell, NAVY)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.name = "Arial"
        row_idx = 1

    # ── Data rows: zebra striping ──
    for i, data_row in enumerate(rows_data):
        for col_idx, text in enumerate(data_row[:num_cols]):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            if i % 2 == 1:
                _set_cell_shading(cell, SLATE_50)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = SLATE_700_RGB
                    run.font.name = "Arial"
        row_idx += 1

    doc.add_paragraph()  # spacing after table
